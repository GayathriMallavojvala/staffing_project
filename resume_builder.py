"""
Resume Builder v2
--------------------
Two outputs from the same data:
  1. render_resume_html()  - a styled HTML preview shown live in the app as the
     user types (Streamlit reruns on every widget change, so this updates the
     same way Overleaf's preview pane updates as you edit).
  2. build_resume_pdf()    - the same design, typeset as a real PDF via
     reportlab, for the actual download. A polished PDF matters more than a
     docx here since PDF is what actually gets submitted to most ATS/clients.

ATS-friendly by construction: single column, no tables/text boxes/images (the
most common reason ATS parsers fail to read a resume at all), standard section
headers matching what ats_scorer.py itself checks for.
"""

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
from reportlab.lib import colors

THEMES = {
    "Navy": {"accent": "#1E3A8A", "accent_rgb": colors.HexColor("#1E3A8A"), "text": "#1F2937"},
    "Charcoal": {"accent": "#111827", "accent_rgb": colors.HexColor("#111827"), "text": "#1F2937"},
    "Teal": {"accent": "#0F766E", "accent_rgb": colors.HexColor("#0F766E"), "text": "#1F2937"},
}


def _section_list(data, key):
    return data.get(key) or []


def render_resume_html(data, theme="Navy"):
    """Live preview HTML, styled to closely match the PDF output."""
    t = THEMES.get(theme, THEMES["Navy"])
    name = data.get("name") or "Your Name"
    contact_bits = [b for b in [data.get("email"), data.get("phone"), data.get("linkedin")] if b]
    contact_line = " &nbsp;|&nbsp; ".join(contact_bits)

    def section_html(title, inner_html):
        if not inner_html:
            return ""
        return (f'<div style="margin-top:16px;">'
                f'<div style="font-size:12px;font-weight:700;letter-spacing:0.06em;color:{t["accent"]};'
                f'text-transform:uppercase;border-bottom:1.5px solid {t["accent"]};padding-bottom:3px;">{title}</div>'
                f'<div style="margin-top:6px;">{inner_html}</div></div>')

    summary_html = f'<div style="font-size:13.5px;color:{t["text"]};">{data.get("summary","")}</div>' if data.get("summary") else ""
    skills_html = f'<div style="font-size:13.5px;color:{t["text"]};">{", ".join(data.get("skills", []))}</div>' if data.get("skills") else ""

    exp_html = ""
    for exp in _section_list(data, "experience"):
        bullets = "".join(f'<li style="margin-bottom:2px;">{b}</li>' for b in exp.get("bullets", []) if b.strip())
        exp_html += (f'<div style="margin-bottom:8px;">'
                     f'<span style="font-weight:700;font-size:13.5px;">{exp.get("title","")} — {exp.get("company","")}</span>'
                     f'<span style="font-size:12px;color:#6B7280;font-style:italic;"> &nbsp;{exp.get("dates","")}</span>'
                     f'<ul style="margin:4px 0 0 18px;padding:0;font-size:13px;color:{t["text"]};">{bullets}</ul></div>')

    proj_html = ""
    for proj in _section_list(data, "projects"):
        bullets = "".join(f'<li style="margin-bottom:2px;">{b}</li>' for b in proj.get("bullets", []) if b.strip())
        proj_html += (f'<div style="margin-bottom:8px;">'
                      f'<span style="font-weight:700;font-size:13.5px;">{proj.get("name","")}</span>'
                      f'<ul style="margin:4px 0 0 18px;padding:0;font-size:13px;color:{t["text"]};">{bullets}</ul></div>')

    edu_html = ""
    for edu in _section_list(data, "education"):
        parts = [edu.get("degree", ""), edu.get("institution", "")]
        line = " — ".join(p for p in parts if p)
        extra = " | ".join(x for x in [edu.get("year", ""), edu.get("score", "")] if x)
        edu_html += f'<div style="font-size:13px;margin-bottom:2px;">{line} {f"({extra})" if extra else ""}</div>'

    certs = _section_list(data, "certifications")
    cert_html = "".join(f'<li style="font-size:13px;margin-bottom:2px;">{c}</li>' for c in certs if c.strip())
    if cert_html:
        cert_html = f'<ul style="margin:0 0 0 18px;padding:0;">{cert_html}</ul>'

    html = f"""
    <div style="background:white; border:1px solid #E5E7EB; border-radius:4px; padding:36px 40px;
                max-width:720px; margin:auto; font-family:'Segoe UI',Arial,sans-serif; min-height:500px;
                box-shadow:0 1px 6px rgba(0,0,0,0.06);">
        <div style="text-align:center;">
            <div style="font-size:24px; font-weight:800; color:{t['accent']};">{name}</div>
            <div style="font-size:12px; color:#6B7280; margin-top:4px;">{contact_line}</div>
        </div>
        {section_html("Summary", summary_html)}
        {section_html("Skills", skills_html)}
        {section_html("Experience", exp_html)}
        {section_html("Projects", proj_html)}
        {section_html("Education", edu_html)}
        {section_html("Certifications", cert_html)}
    </div>
    """
    return html


def build_resume_pdf(data, output_path, theme="Navy"):
    t = THEMES.get(theme, THEMES["Navy"])

    doc = SimpleDocTemplate(output_path, pagesize=LETTER,
                             topMargin=0.6*inch, bottomMargin=0.6*inch,
                             leftMargin=0.75*inch, rightMargin=0.75*inch)

    name_style = ParagraphStyle("Name", fontName="Helvetica-Bold", fontSize=22, leading=26,
                                 textColor=t["accent_rgb"], alignment=TA_CENTER, spaceAfter=6)
    contact_style = ParagraphStyle("Contact", fontName="Helvetica", fontSize=9.5, leading=13,
                                    textColor=colors.HexColor("#6B7280"), alignment=TA_CENTER, spaceAfter=10)
    section_style = ParagraphStyle("Section", fontName="Helvetica-Bold", fontSize=11, leading=14,
                                    textColor=t["accent_rgb"], spaceBefore=12, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontName="Helvetica", fontSize=10, textColor=colors.HexColor("#1F2937"),
                                 leading=14)
    bullet_style = ParagraphStyle("Bullet", fontName="Helvetica", fontSize=10, textColor=colors.HexColor("#1F2937"),
                                   leading=14, leftIndent=12)
    title_style = ParagraphStyle("EntryTitle", fontName="Helvetica-Bold", fontSize=10.5, leading=14,
                                  textColor=colors.HexColor("#111827"), spaceBefore=4)

    story = []
    story.append(Paragraph(data.get("name", "Your Name"), name_style))
    contact_bits = [b for b in [data.get("email"), data.get("phone"), data.get("linkedin")] if b]
    if contact_bits:
        story.append(Paragraph(" &nbsp;|&nbsp; ".join(contact_bits), contact_style))

    def add_section(title):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=1.2, color=t["accent_rgb"], spaceAfter=6))

    if data.get("summary"):
        add_section("Summary")
        story.append(Paragraph(data["summary"], body_style))

    if data.get("skills"):
        add_section("Skills")
        story.append(Paragraph(", ".join(data["skills"]), body_style))

    if data.get("experience"):
        add_section("Experience")
        for exp in data["experience"]:
            story.append(Paragraph(f"{exp.get('title','')} — {exp.get('company','')}"
                                    f"{'  ' + exp.get('dates','') if exp.get('dates') else ''}", title_style))
            bullets = [ListItem(Paragraph(b, bullet_style)) for b in exp.get("bullets", []) if b.strip()]
            if bullets:
                story.append(ListFlowable(bullets, bulletType="bullet", start="•", leftIndent=14))

    if data.get("projects"):
        add_section("Projects")
        for proj in data["projects"]:
            story.append(Paragraph(proj.get("name", ""), title_style))
            bullets = [ListItem(Paragraph(b, bullet_style)) for b in proj.get("bullets", []) if b.strip()]
            if bullets:
                story.append(ListFlowable(bullets, bulletType="bullet", start="•", leftIndent=14))

    if data.get("education"):
        add_section("Education")
        for edu in data["education"]:
            parts = [edu.get("degree", ""), edu.get("institution", "")]
            line = " — ".join(p for p in parts if p)
            extra = " | ".join(x for x in [edu.get("year", ""), edu.get("score", "")] if x)
            if extra:
                line += f" ({extra})"
            story.append(Paragraph(line, body_style))

    if data.get("certifications"):
        add_section("Certifications")
        certs = [ListItem(Paragraph(c, bullet_style)) for c in data["certifications"] if c.strip()]
        if certs:
            story.append(ListFlowable(certs, bulletType="bullet", start="•", leftIndent=14))

    doc.build(story)
    return output_path


def build_resume_docx(data, output_path):
    """Kept as a secondary, editable export format alongside the PDF."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(data.get("name", "Your Name"))
    run.bold = True
    run.font.size = Pt(20)

    contact_bits = [b for b in [data.get("email"), data.get("phone"), data.get("linkedin")] if b]
    if contact_bits:
        contact_p = doc.add_paragraph(" | ".join(contact_bits))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    if data.get("summary"):
        add_section_heading("Summary")
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        add_section_heading("Skills")
        doc.add_paragraph(", ".join(data["skills"]))

    if data.get("experience"):
        add_section_heading("Experience")
        for exp in data["experience"]:
            line = doc.add_paragraph()
            r1 = line.add_run(f"{exp.get('title','')} — {exp.get('company','')}")
            r1.bold = True
            if exp.get("dates"):
                r2 = line.add_run(f"  ({exp['dates']})")
                r2.italic = True
            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style="List Bullet")

    if data.get("projects"):
        add_section_heading("Projects")
        for proj in data["projects"]:
            line = doc.add_paragraph()
            r1 = line.add_run(proj.get("name", ""))
            r1.bold = True
            for bullet in proj.get("bullets", []):
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style="List Bullet")

    if data.get("education"):
        add_section_heading("Education")
        for edu in data["education"]:
            parts = [edu.get("degree", ""), edu.get("institution", "")]
            line = " — ".join(p for p in parts if p)
            if edu.get("year"):
                line += f" ({edu['year']})"
            if edu.get("score"):
                line += f" | {edu['score']}"
            doc.add_paragraph(line)

    if data.get("certifications"):
        add_section_heading("Certifications")
        for cert in data["certifications"]:
            if cert.strip():
                doc.add_paragraph(cert.strip(), style="List Bullet")

    doc.save(output_path)
    return output_path
